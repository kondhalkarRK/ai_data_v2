"""Zero-LLM DOCX extraction for OKF/RAG business knowledge."""
from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone

from features.okf_knowledge.md_extractor import _clean_text, _split_oversized

try:
    from docx import Document

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def docx_extraction_available() -> bool:
    return _DOCX_AVAILABLE


def extract_docx_to_concepts(file_bytes: bytes, source_filename: str) -> list[dict]:
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for DOCX extraction.")

    document = Document(io.BytesIO(file_bytes))
    doc_id = str(uuid.uuid4())[:8]
    ingested_at = datetime.now(timezone.utc).isoformat()
    sections: list[tuple[str, list[str]]] = []
    title = f"{source_filename} — overview"
    body: list[str] = []

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text or "")
        if not text:
            continue
        style = str(getattr(paragraph.style, "name", "") or "").lower()
        if style.startswith("heading"):
            if body:
                sections.append((title, body))
            title, body = text, []
        else:
            body.append(text)
    if body:
        sections.append((title, body))

    for table in document.tables:
        rows = [
            " | ".join(_clean_text(cell.text or "") for cell in row.cells)
            for row in table.rows
        ]
        if rows:
            sections.append((f"{source_filename} — table", rows))

    concepts: list[dict] = []
    for section_number, (section_title, paragraphs) in enumerate(sections, start=1):
        text = _clean_text("\n\n".join(paragraphs))
        if len(text) < 40:
            continue
        for chunk in _split_oversized(text):
            concepts.append(
                {
                    "concept_id": f"{doc_id}_{section_number}_{len(concepts)}",
                    "title": section_title[:90],
                    "source_doc": source_filename,
                    "source_page": section_number,
                    "source_locator": f"section {section_number}",
                    "tags": ["docx", "business-doc"],
                    "body": f"{section_title}\n\n{chunk}"[:1400],
                    "ingested_at": ingested_at,
                    "doc_type": "business_document",
                }
            )
    return concepts


def extract_docx_file(path) -> list[dict]:
    from pathlib import Path

    file_path = Path(path)
    return extract_docx_to_concepts(file_path.read_bytes(), file_path.name)
